from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs") / "JESD_TONE_48bit_DDS原理与计算公式.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], "E8EEF5")
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row_data):
            cells[i].text = text
    return table


def add_formula(doc, text):
    p = doc.add_paragraph()
    p.style = "Formula"
    r = p.add_run(text)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    formula = styles.add_style("Formula", 1)
    formula.font.name = "Consolas"
    formula._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    formula.font.size = Pt(10)
    formula.paragraph_format.left_indent = Inches(0.25)
    formula.paragraph_format.space_before = Pt(2)
    formula.paragraph_format.space_after = Pt(6)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("JESD Tone 48-bit DDS 原理与计算公式")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle.add_run("适用范围：当前 KU5P + AD9173 工程中的 FPGA 内部 JESD single-tone 输出路径。")
    subtitle.paragraph_format.space_after = Pt(14)

    doc.add_heading("1. 核心思想", level=1)
    doc.add_paragraph(
        "本设计中的 JESD single-tone 不再依赖 AD9173 片内 NCO 直接产生正弦，而是在 FPGA 侧使用 "
        "48-bit FTW 相位累加器产生相位，再通过 Vivado DDS Compiler 的 phase-to-amplitude 模块转换成 "
        "signed 16-bit 正弦样点，经 JESD payload 送入 AD9173。"
    )
    add_bullet(doc, "频率控制由 48-bit FTW 决定，频率分辨率来自 48-bit 相位累加器。")
    add_bullet(doc, "幅度样点由 DDS Compiler LUT-only phase-to-sine 模块产生，当前相位输入宽度为 16 bit。")
    add_bullet(doc, "JESD 时钟每拍并行输出 4 个连续样点，因此单路 payload sample rate 是 JESD user clock 的 4 倍。")

    doc.add_heading("2. 相位累加与 FTW", level=1)
    doc.add_paragraph("DDS 的基本状态量是相位累加器 phase。每输出一个样点，相位累加一次：")
    add_formula(doc, "phase[n+1] = (phase[n] + FTW) mod 2^48")
    doc.add_paragraph("输出频率与 FTW 的关系为：")
    add_formula(doc, "f_out = FTW * Fs / 2^48")
    add_formula(doc, "FTW = round(f_out * 2^48 / Fs)")
    doc.add_paragraph("其中 Fs 是单路 DAC payload 样点率，而不是单纯的 JESD user clock。")

    doc.add_heading("3. 当前工程中的采样率", level=1)
    add_formula(doc, "f_clk = 245.76 MHz")
    add_formula(doc, "samples_per_clk = 4")
    add_formula(doc, "Fs = f_clk * 4 = 983.04 MSPS")
    add_formula(doc, "Delta_f = Fs / 2^48 = 983.04e6 / 2^48 ≈ 3.49246 uHz")
    doc.add_paragraph(
        "如果只按 245.76 MHz 时钟计算，分辨率会是约 0.873115 uHz；但当前 RTL 中 FTW 的定义是"
        "每个输出样点的相位步进，因此应使用 983.04 MSPS 计算。"
    )

    add_table(
        doc,
        ["项目", "当前值", "说明"],
        [
            ["JESD user clock", "245.76 MHz", "RTL 时钟域中的 beat clock"],
            ["每拍样点数", "4", "每个 DAC converter 每拍输出 4 个连续 16-bit 样点"],
            ["Payload sample rate", "983.04 MSPS", "FTW 公式中的 Fs"],
            ["相位累加器宽度", "48 bit", "决定频率控制字分辨率"],
            ["DDS Compiler phase input", "16 bit", "当前使用 phase[47:32] 进入 phase-to-sine"],
            ["DDS output", "signed 16 bit sine", "送入后续 scale_sample() 幅度缩放"],
        ],
        [2500, 2300, 4560],
    )

    doc.add_heading("4. 并行 4 样点的相位关系", level=1)
    doc.add_paragraph("由于一个 JESD beat 内需要 4 个连续样点，RTL 在同一拍内展开 4 个相位：")
    add_formula(doc, "phase0 = P")
    add_formula(doc, "phase1 = P + FTW")
    add_formula(doc, "phase2 = P + 2*FTW")
    add_formula(doc, "phase3 = P + 3*FTW")
    add_formula(doc, "P_next = P + 4*FTW")
    doc.add_paragraph(
        "这样虽然 FPGA 时钟是 245.76 MHz，但每个 DAC 通道看到的是连续的 983.04 MSPS 样点流。"
    )

    doc.add_heading("5. Phase-to-Amplitude 转换", level=1)
    doc.add_paragraph(
        "当前 Vivado DDS Compiler 配置为 SIN_COS_LUT_only、Phase_Width=16、Output_Width=16、Output_Selection=Sine。"
        "RTL 将 48-bit phase 的高 16 bit 送入 DDS Compiler："
    )
    add_formula(doc, "phase_word = phase[47:32]")
    doc.add_paragraph(
        "这意味着频率控制仍保留 48-bit FTW 的细分能力；但正弦幅度查表使用 16-bit 相位地址。"
        "低 32 bit 主要用于长期相位累加和频率分辨率，高 16 bit 决定每个样点进入正弦表的位置。"
    )

    doc.add_heading("6. 幅度缩放关系", level=1)
    doc.add_paragraph("DDS Compiler 输出 signed 16-bit 正弦样点 raw_sample，后级使用 Q1.15 风格的 scale 缩放：")
    add_formula(doc, "scaled_sample ≈ raw_sample * scale / 2^15")
    add_bullet(doc, "scale = 0 表示静音。")
    add_bullet(doc, "scale = 0x7fff 表示 JESD 数字路径接近满幅。")
    add_bullet(doc, "JESD single-tone 当前不套用 NCO/RAM 幅度校准表时，3 Vpp 这类超过 dac_full_scale_vpk 的目标会被数字侧夹到 0x7fff。")

    doc.add_heading("7. 频率覆盖与示例", level=1)
    doc.add_paragraph("理论上，只要模拟链路和 AD9173/JESD 设置允许，DDS 频率由 FTW 直接指定：")
    add_formula(doc, "FTW(1 mHz) = round(1e-3 * 2^48 / 983.04e6) ≈ 286")
    add_formula(doc, "FTW(200 MHz) = round(200e6 * 2^48 / 983.04e6) ≈ 0x341555555555")
    doc.add_paragraph(
        "1 mHz 到 200 MHz 在 48-bit FTW 数值上都可表示。实际输出质量还取决于时钟相噪、DDS 相位截断杂散、"
        "DAC 与模拟链路带宽、输出耦合方式、幅度校准和测试时间窗口。"
    )

    doc.add_heading("8. 与片内 NCO / RAM 模式的边界", level=1)
    add_bullet(doc, "JESD single-tone：FPGA 侧 DDS 产生样点，经 JESD 送入 AD9173。")
    add_bullet(doc, "NCO-only：HostApp/VIO/SPI 配置 AD9173 片内 NCO；此路径不由 FPGA DDS 样点决定。")
    add_bullet(doc, "RAM waveform：HostApp 写入 waveform RAM，RTL 循环播放 RAM 样点；不是本 DDS 单音路径。")
    doc.add_paragraph(
        "因此，JESD single-tone 的 FTW、scale 和校准策略应与 NCO-only、RAM waveform 分开看待。"
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
